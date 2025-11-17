from __future__ import annotations
import json
from typing import Optional
from io import BytesIO
from datetime import datetime
import pdfplumber
from docx import Document
from app.services.llm import get_openai
from app.settings import settings
from app.models import CandidateCVNormalized
from app.services.cv_parser import parse_cv_bytes_to_normalized

CV_PARSER_SYSTEM = """You are a recruitment assistant that extracts structured candidate CV data from free-form text.
Extract the following information accurately:
- Candidate identity: full_name (required), email, phone, linkedin_url, location (city, region, country, remote_preference), right_to_work, notice_period_weeks, availability_date, current_compensation, target_compensation
- Experience: List of work experience items with title, employer, location, start_date (YYYY-MM format), end_date (YYYY-MM or null if current), is_current, responsibilities, achievements, technologies, team_size
- Skills: List of skills with name, category (tech/sales/analytics/management/language/other), level (novice/intermediate/advanced/expert), evidence
- Education: List of education items with institution, degree, field, start_year, end_year
- Certifications: List of certification names
- Languages: List of language proficiencies with name and proficiency (basic/conversational/fluent/native)
- Documents: resume_url, cover_letter_url

Be accurate and conservative. If information is not provided in the text, use null for optional fields.
For dates, use YYYY-MM format (e.g., "2022-01").
For experience items, set is_current=true if the end_date is null or if the text indicates current employment.
For skills, infer category and level from context if not explicitly stated.
Output valid JSON matching the CandidateCVNormalized schema."""


def _extract_text_from_bytes(data: bytes, filename: Optional[str] = None) -> str:
    """
    Extract text from PDF, DOCX, or plain text file.
    Returns the extracted text as a string.
    """
    if filename:
        ext = filename.lower().split('.')[-1]
    else:
        # Try to detect from content
        if data.startswith(b'%PDF'):
            ext = 'pdf'
        elif data.startswith(b'PK\x03\x04'):  # ZIP/DOCX signature
            ext = 'docx'
        else:
            ext = 'txt'
    
    if ext == 'pdf':
        try:
            with pdfplumber.open(BytesIO(data)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    elif ext == 'docx':
        try:
            doc = Document(BytesIO(data))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    else:
        # Assume plain text
        try:
            return data.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return data.decode('latin-1')
            except Exception as e:
                raise ValueError(f"Failed to decode text file: {e}")


def parse_cv_bytes_to_normalized_llm(
    data: bytes,
    filename: Optional[str] = None,
    parser_version: str = "cvx-1.2.0"
) -> CandidateCVNormalized:
    """
    Parse CV file (PDF/DOCX/TXT) using LLM extraction.
    
    First extracts text from the file, then uses LLM to extract structured data
    matching the CandidateCVNormalized schema.
    
    Falls back to stub parser if OpenAI API key is not configured.
    """
    # Try to get OpenAI client, fallback to stub if not configured
    try:
        client_openai = get_openai()
    except RuntimeError:
        # API key not configured, use stub fallback
        return parse_cv_bytes_to_normalized(data, filename=filename, parser_version=parser_version)
    
    # Extract text from file
    try:
        cv_text = _extract_text_from_bytes(data, filename)
    except Exception as e:
        # If text extraction fails, fallback to stub
        print(f"Warning: Failed to extract text from CV file: {e}. Falling back to stub parser.")
        return parse_cv_bytes_to_normalized(data, filename=filename, parser_version=parser_version)
    
    if not cv_text or not cv_text.strip():
        # Empty text, fallback to stub
        print("Warning: CV file appears to be empty. Falling back to stub parser.")
        return parse_cv_bytes_to_normalized(data, filename=filename, parser_version=parser_version)
    
    # Build the user prompt
    user_prompt = f"""Extract structured candidate data from the following CV text:

{cv_text}

Extract and return a valid JSON object matching this structure:
{{
  "candidate": {{
    "full_name": "string (required)",
    "email": "string or null",
    "phone": "string or null",
    "linkedin_url": "string or null",
    "location": {{"city": "string or null", "region": "string or null", "country": "string or null", "remote_preference": "remote|hybrid|onsite|unspecified or null"}} or null,
    "right_to_work": ["string"] or null,
    "notice_period_weeks": number or null,
    "availability_date": "YYYY-MM-DD" or null,
    "current_compensation": {{"base_amount": number or null, "currency": "string (3 chars) or null", "period": "year|month or null", "bonus_ote": number or null, "equity": "string or null"}} or null,
    "target_compensation": {{"base_min": number or null, "base_max": number or null, "currency": "string (3 chars) or null", "period": "year|month or null"}} or null
  }},
  "experience": [
    {{
      "title": "string (required)",
      "employer": "string (required)",
      "location": "string or null",
      "start_date": "YYYY-MM (required)",
      "end_date": "YYYY-MM or null",
      "is_current": boolean,
      "responsibilities": ["string"] or null,
      "achievements": ["string"] or null,
      "technologies": ["string"] or null,
      "team_size": number or null
    }}
  ],
  "skills": [
    {{
      "name": "string (required)",
      "category": "tech|sales|analytics|management|language|other or null",
      "level": "novice|intermediate|advanced|expert or null",
      "evidence": ["string"] or null
    }}
  ],
  "education": [
    {{
      "institution": "string or null",
      "degree": "string or null",
      "field": "string or null",
      "start_year": number or null,
      "end_year": number or null
    }}
  ] or null,
  "certifications": ["string"] or null,
  "languages": [
    {{
      "name": "string (required)",
      "proficiency": "basic|conversational|fluent|native or null"
    }}
  ] or null,
  "documents": {{
    "resume_url": "string or null",
    "cover_letter_url": "string or null"
  }} or null,
  "extraction_meta": {{
    "source": "pdf|docx|linkedin|text|other",
    "extracted_at": "ISO 8601 datetime",
    "parser_version": "string"
  }} or null
}}

Return only valid JSON, no markdown formatting or explanation."""

    try:
        # Use JSON mode for structured output
        resp = client_openai.chat.completions.create(
            model=settings.openai_model_long,  # Use long model for better extraction quality
            messages=[
                {"role": "system", "content": CV_PARSER_SYSTEM},
                {"role": "user", "content": user_prompt},
            ],
            response_format={"type": "json_object"},  # Force JSON output
            temperature=0.1,  # Low temperature for consistent, factual extraction
            max_tokens=4000,  # Enough for structured CV data
        )
        
        # Extract JSON from response
        content = resp.choices[0].message.content.strip()
        
        # Remove markdown code blocks if present
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        # Parse JSON
        cv_data = json.loads(content)
        
        # Ensure extraction_meta is set
        if "extraction_meta" not in cv_data or cv_data["extraction_meta"] is None:
            source_type = "pdf" if filename and filename.lower().endswith(".pdf") else \
                         "docx" if filename and filename.lower().endswith(".docx") else "text"
            cv_data["extraction_meta"] = {
                "source": source_type,
                "extracted_at": datetime.utcnow().isoformat(),
                "parser_version": parser_version
            }
        else:
            # Update parser version and timestamp
            cv_data["extraction_meta"]["parser_version"] = parser_version
            if not cv_data["extraction_meta"].get("extracted_at"):
                cv_data["extraction_meta"]["extracted_at"] = datetime.utcnow().isoformat()
        
        # Validate and create Pydantic model
        return CandidateCVNormalized.model_validate(cv_data)
        
    except json.JSONDecodeError as e:
        # If JSON parsing fails, fallback to stub
        print(f"Warning: LLM returned invalid JSON: {e}. Falling back to stub parser.")
        return parse_cv_bytes_to_normalized(data, filename=filename, parser_version=parser_version)
    except Exception as e:
        # If LLM call fails, fallback to stub
        print(f"Warning: CV parsing LLM call failed: {e}. Falling back to stub parser.")
        return parse_cv_bytes_to_normalized(data, filename=filename, parser_version=parser_version)

